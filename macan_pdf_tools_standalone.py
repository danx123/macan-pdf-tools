#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Macan PDF Tools — Standalone Edition
=====================================
Extracted from macan_converter (PDF Tools page) into a standalone application.

Specifically designed for low-spec PCs/laptops and older CPUs that DO NOT support AVX/AVX2 instructions (e.g., early-generation Core i3/i5, Celeron, Atom, and some older AMD processors). Therefore:

- DOES NOT use numpy / opencv (the official opencv-python and numpy versions are often compiled with AVX2 and will crash with "Illegal instruction"
on non-AVX CPUs).
- All image processing uses pure Pillow (PIL).
- PDF page rendering uses pypdfium2 (binding to PDFium, no AVX required).
- PDF structure manipulation (merge & stream compression) uses pikepdf (optional,
Automatic fallback to pure pypdfium2 if pikepdf is not installed).

Features (4 sub-tools, similar to the "PDF Tools" page in Macan Converter):
1. Image to PDF — merge images into one or multiple PDFs
2. PDF to Image — export PDF pages to PNG/JPG/WEBP
3. PDF Merger — merge multiple PDFs with optional compression
4. PDF Document Conversion — PDF -> TXT / PDF -> DOCX / PDF -> XLSX

Supports 2 languages: Indonesian & English (select from the top right corner).

Dependencies (all lightweight, no AVX requirement):
pip install PySide6 Pillow pypdfium2 pikepdf python-docx openpyxl --break-system-packages

Note: pikepdf, python-docx, and openpyxl are optional — the related features will be automatically disabled (with a clear message) if the libraries are not installed.
The application will continue to run normally for other features.
"""

import sys
# Menipu sistem agar mengira NumPy sudah gagal diimpor
sys.modules['numpy'] = None  
import os
import io
import json
import subprocess
import urllib.request
import urllib.error

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QGridLayout,
    QLabel, QPushButton, QListWidget, QListWidgetItem, QAbstractItemView,
    QFileDialog, QLineEdit, QComboBox, QSpinBox, QFrame, QProgressBar,
    QMessageBox, QStackedWidget, QSplitter, QScrollArea, QSizePolicy
)
from PySide6.QtCore import Qt, QSize, QThread, QObject, Signal, Slot, QRunnable, QThreadPool, QSettings, QByteArray, QUrl
from PySide6.QtGui import QIcon, QPixmap, QPainter, QDragEnterEvent, QDragMoveEvent, QDropEvent, QDesktopServices, QFont
from PySide6.QtSvg import QSvgRenderer

from PIL import Image, ImageOps

try:
    import pypdfium2 as pdfium
    HAS_PDFIUM = True
except ImportError:
    HAS_PDFIUM = False

try:
    import pikepdf
    HAS_PIKEPDF = True
except ImportError:
    HAS_PIKEPDF = False

try:
    import docx as _docx_check  # python-docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl as _openpyxl_check
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


APP_VERSION = "1.2.0"
ORG_NAME = "MacanAngkasa"
APP_NAME = "MacanPdfToolsStandalone"
IMAGE_EXT = ['.png', '.jpg', '.jpeg', '.bmp', '.webp', '.gif']
PDF_EXT = ['.pdf']

UPDATE_JSON_URL = "https://raw.githubusercontent.com/danx123/macan-pdf-tools/main/version.json"
REPO_URL = "https://github.com/danx123/macan-pdf-tools"


def get_app_dir():
    """Folder tempat script/exe ini berada (aman untuk dev & hasil Nuitka/PyInstaller)."""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


LOGO_PATH = os.path.join(get_app_dir(), "logo.png")


# ──────────────────────────────────────────────────────────────────────────
#  Bahasa / Language strings
# ──────────────────────────────────────────────────────────────────────────
LANGUAGES = {
    "id": {
        "window_title": "Macan PDF Tools — Standalone (v{version})",
        "nav_img2pdf": "Image to PDF",
        "nav_pdf2img": "PDF to Image",
        "nav_merger": "PDF Merger",
        "nav_docconv": "Konversi Dokumen PDF",
        "lang_label": "Bahasa:",
        "add_files_btn": "Tambah File",
        "clear_btn": "Bersihkan",
        "output_label": "Output:",
        "output_placeholder": "Pilih folder output...",
        "choose_folder_btn": "Pilih Folder",
        "choose_folder_title": "Pilih Folder Output",
        "status_ready": "Siap.",
        "status_stopped": "Dihentikan oleh pengguna.",
        "start_btn": "Mulai",
        "stop_btn": "Stop",
        "options_label": "Opsi",
        "drop_placeholder": "Seret file ke sini atau klik 'Tambah File'",
        "error_title": "Error",
        "done_title": "Selesai",
        "error_no_output": "Pilih folder output terlebih dahulu.",
        "dep_missing_title": "Dependency hilang",
        "dep_missing_body": "Beberapa fitur tidak akan berfungsi:\n\n{items}\n\nInstall dengan:\npip install {pkgs} --break-system-packages",

        # Image to PDF
        "img2pdf_title": "Image to PDF",
        "img2pdf_start_btn": "Gabungkan -> PDF",
        "img2pdf_no_files": "Tambahkan file gambar terlebih dahulu.",
        "quality_label": "Kualitas:",
        "qualities": ["Maksimum (100)", "Bagus (95)", "Baik (85)", "Sedang (75)", "Rendah (50)"],
        "out_mode_label": "Output:",
        "out_modes": ["Satu PDF untuk semua gambar", "Satu PDF per gambar"],
        "target_size_label": "Target Ukuran:",
        "target_size_hint": "0 = otomatis",

        # PDF to Image
        "pdf2img_title": "PDF to Image",
        "pdf2img_start_btn": "Konversi -> Gambar",
        "pdf2img_no_files": "Tambahkan file PDF terlebih dahulu.",
        "format_label": "Format:",
        "dpi_label": "Kualitas (DPI):",
        "dpis": ["Rendah (72 DPI)", "Sedang (150 DPI)", "Baik (200 DPI)", "Tinggi (300 DPI)", "Maksimum (600 DPI)"],
        "pages_label": "Halaman:",
        "pages_hint": "Cth: 1,3,5-8 (kosong = semua)",

        # Merger
        "merger_title": "PDF Merger",
        "merger_start_btn": "Gabungkan PDF",
        "merger_no_files": "Tambahkan minimal 2 file PDF.",
        "merger_qualities": ["Pertahankan asli", "Optimal (85)", "Rendah (60)"],
        "merger_pikepdf_note": "Catatan: 'pikepdf' tidak terpasang — kompresi tidak\ntersedia, hanya merge struktural (via pypdfium2).",

        # Doc conversion
        "docconv_start_btn": "Konversi",
        "docconv_no_files": "Tambahkan file PDF terlebih dahulu.",
        "convert_to_label": "Konversi ke:",
        "missing_libs_label": "Belum terpasang: {items}",

        # Estimated size
        "estimated_size_label": "Estimasi Ukuran:",
        "estimated_size_empty": "-",
        "estimated_size_calc": "± {size}",
        "estimated_size_target": "≤ {target} (target)",

        # Open output folder after done
        "open_folder_failed": "Gagal membuka folder output:\n{err}",

        # About page
        "nav_about": "Tentang",
        "about_app_title": "Macan PDF Tools",
        "about_edition": "Standalone Edition",
        "about_version_label": "Versi:",
        "about_tagline": "Toolkit PDF ringan & cepat — dioptimalkan untuk PC/laptop low-spec tanpa AVX/AVX2.",
        "about_description": (
            "Macan PDF Tools adalah aplikasi desktop ringan untuk mengelola dokumen PDF "
            "sehari-hari: mengubah gambar menjadi PDF, mengekstrak halaman PDF menjadi gambar, "
            "menggabungkan beberapa PDF sekaligus mengompresinya, serta mengonversi PDF ke "
            "format TXT, DOCX, dan XLSX. Seluruh pemrosesan gambar memakai Pillow murni dan "
            "rendering PDF memakai pypdfium2, sehingga aman dijalankan pada CPU lama yang tidak "
            "mendukung instruksi AVX/AVX2."
        ),
        "about_features_title": "Fitur Utama",
        "about_features": [
            "Image to PDF — gabungkan banyak gambar menjadi satu atau beberapa file PDF",
            "PDF to Image — ekspor halaman PDF ke PNG / JPG / WEBP dengan kontrol DPI",
            "PDF Merger — gabungkan PDF dengan opsi kompresi ukuran target",
            "Konversi Dokumen — PDF ke TXT, DOCX, dan XLSX",
        ],
        "about_dependencies_title": "Status Dependency",
        "about_dep_required": "(wajib)",
        "about_dep_optional": "(opsional)",
        "about_dep_installed": "Terpasang",
        "about_dep_missing": "Tidak terpasang",
        "about_update_title": "Pembaruan Aplikasi",
        "about_check_update_btn": "Cek Pembaruan",
        "about_checking_update": "Memeriksa pembaruan...",
        "about_up_to_date": "Aplikasi sudah versi terbaru ({version}).",
        "about_update_available": "Versi baru tersedia: {version} (saat ini {current}).",
        "about_update_error": "Gagal memeriksa pembaruan:\n{err}",
        "about_update_notes_title": "Catatan rilis:",
        "about_open_download_btn": "Buka Halaman Unduhan",
        "about_open_repo_btn": "Kunjungi Repository",
        "about_credits": "Dikembangkan oleh Macan Angkasa. Dibangun dengan Python, PySide6, Pillow, dan pypdfium2.",
        "about_copyright": "© {year} Macan Angkasa — Hak cipta dilindungi.",
    },
    "en": {
        "window_title": "Macan PDF Tools — Standalone (v{version})",
        "nav_img2pdf": "Image to PDF",
        "nav_pdf2img": "PDF to Image",
        "nav_merger": "PDF Merger",
        "nav_docconv": "PDF Document Conversion",
        "lang_label": "Language:",
        "add_files_btn": "Add Files",
        "clear_btn": "Clear",
        "output_label": "Output:",
        "output_placeholder": "Select output folder...",
        "choose_folder_btn": "Choose Folder",
        "choose_folder_title": "Select Output Folder",
        "status_ready": "Ready.",
        "status_stopped": "Stopped by user.",
        "start_btn": "Start",
        "stop_btn": "Stop",
        "options_label": "Options",
        "drop_placeholder": "Drag files here or click 'Add Files'",
        "error_title": "Error",
        "done_title": "Done",
        "error_no_output": "Please select an output folder first.",
        "dep_missing_title": "Missing dependency",
        "dep_missing_body": "Some features will not work:\n\n{items}\n\nInstall with:\npip install {pkgs} --break-system-packages",

        # Image to PDF
        "img2pdf_title": "Image to PDF",
        "img2pdf_start_btn": "Combine -> PDF",
        "img2pdf_no_files": "Please add image files first.",
        "quality_label": "Quality:",
        "qualities": ["Maximum (100)", "Good (95)", "Good (85)", "Medium (75)", "Low (50)"],
        "out_mode_label": "Output:",
        "out_modes": ["Single PDF for all images", "One PDF per image"],
        "target_size_label": "Target Size:",
        "target_size_hint": "0 = automatic",

        # PDF to Image
        "pdf2img_title": "PDF to Image",
        "pdf2img_start_btn": "Convert -> Images",
        "pdf2img_no_files": "Please add PDF files first.",
        "format_label": "Format:",
        "dpi_label": "Quality (DPI):",
        "dpis": ["Low (72 DPI)", "Medium (150 DPI)", "Good (200 DPI)", "High (300 DPI)", "Maximum (600 DPI)"],
        "pages_label": "Pages:",
        "pages_hint": "E.g.: 1,3,5-8 (blank = all)",

        # Merger
        "merger_title": "PDF Merger",
        "merger_start_btn": "Merge PDFs",
        "merger_no_files": "Please add at least 2 PDF files.",
        "merger_qualities": ["Keep original", "Optimized (85)", "Low (60)"],
        "merger_pikepdf_note": "Note: 'pikepdf' is not installed — compression is\nunavailable, only structural merge (via pypdfium2).",

        # Doc conversion
        "docconv_start_btn": "Convert",
        "docconv_no_files": "Please add PDF files first.",
        "convert_to_label": "Convert to:",
        "missing_libs_label": "Not installed: {items}",

        # Estimated size
        "estimated_size_label": "Estimated Size:",
        "estimated_size_empty": "-",
        "estimated_size_calc": "± {size}",
        "estimated_size_target": "≤ {target} (target)",

        # Open output folder after done
        "open_folder_failed": "Failed to open output folder:\n{err}",

        # About page
        "nav_about": "About",
        "about_app_title": "Macan PDF Tools",
        "about_edition": "Standalone Edition",
        "about_version_label": "Version:",
        "about_tagline": "A lightweight, fast PDF toolkit — optimized for low-spec PCs/laptops without AVX/AVX2.",
        "about_description": (
            "Macan PDF Tools is a lightweight desktop application for everyday PDF management: "
            "turning images into PDFs, exporting PDF pages as images, merging multiple PDFs while "
            "compressing them, and converting PDFs into TXT, DOCX, and XLSX formats. All image "
            "processing uses pure Pillow and PDF rendering uses pypdfium2, making it safe to run "
            "on older CPUs that do not support AVX/AVX2 instructions."
        ),
        "about_features_title": "Key Features",
        "about_features": [
            "Image to PDF — combine multiple images into one or several PDF files",
            "PDF to Image — export PDF pages to PNG / JPG / WEBP with DPI control",
            "PDF Merger — merge PDFs with optional target-size compression",
            "Document Conversion — PDF to TXT, DOCX, and XLSX",
        ],
        "about_dependencies_title": "Dependency Status",
        "about_dep_required": "(required)",
        "about_dep_optional": "(optional)",
        "about_dep_installed": "Installed",
        "about_dep_missing": "Not installed",
        "about_update_title": "Application Update",
        "about_check_update_btn": "Check for Updates",
        "about_checking_update": "Checking for updates...",
        "about_up_to_date": "You're on the latest version ({version}).",
        "about_update_available": "A new version is available: {version} (current {current}).",
        "about_update_error": "Failed to check for updates:\n{err}",
        "about_update_notes_title": "Release notes:",
        "about_open_download_btn": "Open Download Page",
        "about_open_repo_btn": "Visit Repository",
        "about_credits": "Developed by Macan Angkasa. Built with Python, PySide6, Pillow, and pypdfium2.",
        "about_copyright": "© {year} Macan Angkasa — All rights reserved.",
    },
}


# ──────────────────────────────────────────────────────────────────────────
#  SVG icons — kategori navigasi sidebar (vector, ringan, scalable, no AVX)
# ──────────────────────────────────────────────────────────────────────────
NAV_SVG_ICONS = {
    # Image to PDF — gambar dengan panah ke dokumen
    "img2pdf": """
        <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="none"
             stroke="{color}" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
            <rect x="2" y="4" width="11" height="9" rx="1.2"/>
            <circle cx="5.5" cy="7.5" r="1.1"/>
            <path d="M2.8 12.2l3-3.2 2 2 2.7-3 2.5 3"/>
            <path d="M15.5 8.5h3.5a1.5 1.5 0 0 1 1.5 1.5v9a1.5 1.5 0 0 1-1.5 1.5h-6a1.5 1.5 0 0 1-1.5-1.5v-3"/>
            <path d="M16 12.5h4M16 15h4M16 17.5h2.5"/>
        </svg>""",

    # PDF to Image — dokumen dengan panah ke gambar
    "pdf2img": """
        <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="none"
             stroke="{color}" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
            <path d="M5 2.5h6.5L15 6v9.5a1.5 1.5 0 0 1-1.5 1.5h-7A1.5 1.5 0 0 1 5 15.5z"/>
            <path d="M11.5 2.5V6H15"/>
            <path d="M6.7 9.8h5.5M6.7 12h4"/>
            <rect x="10.5" y="13.5" width="11" height="8" rx="1.2"/>
            <circle cx="13.6" cy="16.4" r="1"/>
            <path d="M11.3 20.5l2.6-2.7 1.7 1.7 2.3-2.5 2.6 2.5"/>
        </svg>""",

    # PDF Merger — dua dokumen menyatu jadi satu
    "merger": """
        <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="none"
             stroke="{color}" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
            <path d="M3 3h6l3 3v9a1.2 1.2 0 0 1-1.2 1.2H3.2A1.2 1.2 0 0 1 2 15V4.2A1.2 1.2 0 0 1 3 3z"/>
            <path d="M9 3v3h3"/>
            <path d="M13 8h6l3 3v9a1.2 1.2 0 0 1-1.2 1.2h-7.6A1.2 1.2 0 0 1 12 20V9.2A1.2 1.2 0 0 1 13.2 8z"/>
            <path d="M19 8v3h3"/>
            <path d="M10.5 12.3l3 3-3 3" stroke-width="2"/>
        </svg>""",

    # PDF Document Conversion — dokumen dengan panah refresh/exchange
    "docconv": """
        <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="none"
             stroke="{color}" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
            <path d="M5 2.5h6.5L15 6v15.5a1.5 1.5 0 0 1-1.5 1.5h-7A1.5 1.5 0 0 1 5 21.5z"/>
            <path d="M11.5 2.5V6H15"/>
            <path d="M6.8 13h6.4M6.8 15.3h6.4M6.8 17.6h4"/>
            <path d="M16 9.8a3.6 3.6 0 0 1 5.8-1.1M22 7.5v2.4h-2.4"/>
            <path d="M22.2 12.4a3.6 3.6 0 0 1-5.8 1.1M16 14.7v-2.4h2.4"/>
        </svg>""",

    # About — info circle
    "about": """
        <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="none"
             stroke="{color}" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
            <circle cx="12" cy="12" r="9.5"/>
            <line x1="12" y1="10.7" x2="12" y2="16.5"/>
            <circle cx="12" cy="7.6" r="1" fill="{color}" stroke="none"/>
        </svg>""",
}


def svg_to_icon(svg_key, color="#C8C8C8", size=20):
    """Render salah satu NAV_SVG_ICONS jadi QIcon. Pure Qt SVG rendering
    (QtSvg, sudah satu paket dengan PySide6) — tidak butuh library tambahan."""
    svg_str = NAV_SVG_ICONS[svg_key].format(color=color)
    pix = QPixmap(size, size)
    pix.fill(Qt.GlobalColor.transparent)
    renderer = QSvgRenderer(QByteArray(svg_str.encode("utf-8")))
    painter = QPainter(pix)
    renderer.render(painter)
    painter.end()
    return QIcon(pix)


def human_size(num_bytes):
    """Format angka byte jadi string ukuran yang manusiawi (KB/MB/GB)."""
    if num_bytes is None:
        return "-"
    size = float(num_bytes)
    for unit in ("B", "KB", "MB", "GB"):
        if size < 1024.0 or unit == "GB":
            return f"{size:.1f} {unit}" if unit != "B" else f"{int(size)} {unit}"
        size /= 1024.0
    return f"{size:.1f} GB"


def open_in_file_manager(folder_path):
    """Buka folder output di file manager OS (Explorer/Finder/Nautilus dkk)."""
    if not folder_path or not os.path.isdir(folder_path):
        return
    try:
        if sys.platform.startswith("win"):
            os.startfile(folder_path)  # noqa: P201 (Windows only)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", folder_path])
        else:
            subprocess.Popen(["xdg-open", folder_path])
    except Exception:
        QDesktopServices.openUrl(QUrl.fromLocalFile(folder_path))


# ──────────────────────────────────────────────────────────────────────────
#  Icon helpers — generic fallback glyph (no numpy/opencv, full Pillow)
# ──────────────────────────────────────────────────────────────────────────
def make_generic_icon(ext, size=96):
    """Bikin icon kotak sederhana berdasarkan ekstensi file. Dipakai sebagai
    placeholder sementara sebelum thumbnail asli selesai dirender, dan
    sebagai fallback kalau file gagal dibaca/di-render."""
    ext = ext.lower().lstrip('.')
    colors = {
        'pdf': (200, 70, 60), 'png': (90, 150, 200), 'jpg': (90, 150, 200),
        'jpeg': (90, 150, 200), 'bmp': (90, 150, 200), 'webp': (90, 150, 200),
        'gif': (90, 150, 200),
    }
    color = colors.get(ext, (120, 120, 120))
    img = Image.new('RGB', (size, size), (45, 45, 45))
    pad = 10
    box = Image.new('RGB', (size - pad * 2, size - pad * 2), color)
    img.paste(box, (pad, pad))
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    pix = QPixmap()
    pix.loadFromData(buf.getvalue(), 'PNG')
    return pix


def _square_thumbnail_pixmap(pil_img, size=96, bg=(45, 45, 45)):
    """Cocokkan gambar PIL apa pun ke kanvas persegi (letterbox), full Pillow,
    tanpa numpy/opencv — aman untuk CPU non-AVX."""
    pil_img = pil_img.convert('RGB')
    fitted = ImageOps.contain(pil_img, (size, size))
    canvas = Image.new('RGB', (size, size), bg)
    off = ((size - fitted.width) // 2, (size - fitted.height) // 2)
    canvas.paste(fitted, off)
    buf = io.BytesIO()
    canvas.save(buf, format='PNG')
    pix = QPixmap()
    pix.loadFromData(buf.getvalue(), 'PNG')
    return pix


# ──────────────────────────────────────────────────────────────────────────
#  Async thumbnail rendering — pure PIL (images) / pypdfium2 (first PDF page)
# ──────────────────────────────────────────────────────────────────────────
class _ThumbnailSignals(QObject):
    ready = Signal(str, QPixmap)


class ThumbnailWorker(QRunnable):
    """Render satu thumbnail di background thread. Tidak memakai numpy/opencv
    sama sekali — gambar lewat Pillow, halaman PDF lewat pypdfium2 (keduanya
    binding ringan tanpa requirement AVX), sehingga aman untuk CPU lama."""

    def __init__(self, file_path, size=96):
        super().__init__()
        self.file_path = file_path
        self.size = size
        self.signals = _ThumbnailSignals()

    def run(self):
        ext = os.path.splitext(self.file_path)[1].lower()
        pix = None
        try:
            if ext == '.pdf':
                if HAS_PDFIUM:
                    pdf = pdfium.PdfDocument(self.file_path)
                    if len(pdf) > 0:
                        page = pdf[0]
                        # render kecil langsung (scale rendah) biar ringan di CPU lemah
                        target_px = max(self.size, 150)
                        w_pt, h_pt = page.get_size()
                        scale = target_px / max(w_pt, h_pt, 1)
                        pil_img = page.render(scale=scale).to_pil()
                        pix = _square_thumbnail_pixmap(pil_img, self.size)
            else:
                pil_img = Image.open(self.file_path)
                pil_img.thumbnail((self.size * 2, self.size * 2))  # decode hemat memori dulu
                pix = _square_thumbnail_pixmap(pil_img, self.size)
        except Exception:
            pix = None

        if pix is None:
            pix = make_generic_icon(ext, self.size)

        self.signals.ready.emit(self.file_path, pix)


# ──────────────────────────────────────────────────────────────────────────
#  Drag & drop file list — async real thumbnails via QThreadPool
# ──────────────────────────────────────────────────────────────────────────
class FileDropArea(QListWidget):
    files_changed = Signal()

    def __init__(self, accept_types=None, lang=None, parent=None):
        super().__init__(parent)
        self.accept_types = accept_types or ['image', 'pdf']
        self.file_paths = []
        self.lang = lang

        self.setAcceptDrops(True)
        self.setDragDropMode(QAbstractItemView.DragDropMode.NoDragDrop)
        self.setViewMode(QListWidget.ViewMode.IconMode)
        self.setIconSize(QSize(96, 96))
        self.setResizeMode(QListWidget.ResizeMode.Adjust)
        self.setWordWrap(True)
        self.setSpacing(10)
        self._icon_cache = {}

        # Batasi jumlah thread biar tetap ringan di CPU low-spec / sedikit core
        self.thread_pool = QThreadPool()
        max_threads = max(1, min(4, QThreadPool.globalInstance().maxThreadCount() // 2))
        self.thread_pool.setMaxThreadCount(max_threads)

        self._placeholder_item = None
        self._set_placeholder()

    def _set_placeholder(self):
        text = self.lang["drop_placeholder"] if self.lang else "Drag files here or click 'Add Files'"
        self._placeholder_item = QListWidgetItem(text)
        self._placeholder_item.setFlags(Qt.ItemFlag.NoItemFlags)
        self._placeholder_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        self.addItem(self._placeholder_item)

    def retranslate(self, lang):
        self.lang = lang
        if self.count() == 1 and self.item(0) == self._placeholder_item:
            self._placeholder_item.setText(lang["drop_placeholder"])

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()

    def dragMoveEvent(self, event: QDragMoveEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()

    def dropEvent(self, event: QDropEvent):
        if event.mimeData().hasUrls():
            paths = [url.toLocalFile() for url in event.mimeData().urls()]
            self.add_files(paths)
            event.acceptProposedAction()
        else:
            event.ignore()

    def _allowed_ext(self):
        allowed = []
        if 'image' in self.accept_types:
            allowed.extend(IMAGE_EXT)
        if 'pdf' in self.accept_types:
            allowed.extend(PDF_EXT)
        return allowed

    def add_files(self, paths):
        if self.count() == 1 and self.item(0) == self._placeholder_item:
            self.clear()
            self.file_paths.clear()

        allowed = self._allowed_ext()
        valid = [p for p in paths if os.path.splitext(p)[1].lower() in allowed and os.path.isfile(p)]

        for path in valid:
            if path in self.file_paths:
                continue
            self.file_paths.append(path)
            ext = os.path.splitext(path)[1].lower().lstrip('.')
            if ext not in self._icon_cache:
                self._icon_cache[ext] = make_generic_icon(ext)
            item = QListWidgetItem(QIcon(self._icon_cache[ext]), os.path.basename(path))
            item.setData(Qt.ItemDataRole.UserRole, path)
            item.setSizeHint(QSize(120, 130))
            item.setTextAlignment(Qt.AlignmentFlag.AlignBottom | Qt.AlignmentFlag.AlignHCenter)
            self.addItem(item)

            # Render thumbnail asli di background — tidak memblokir UI
            worker = ThumbnailWorker(path, size=96)
            worker.signals.ready.connect(self._on_thumbnail_ready)
            self.thread_pool.start(worker)

        if valid:
            self.files_changed.emit()

    @Slot(str, QPixmap)
    def _on_thumbnail_ready(self, file_path, pixmap):
        for i in range(self.count()):
            item = self.item(i)
            if item is self._placeholder_item:
                continue
            if item.data(Qt.ItemDataRole.UserRole) == file_path:
                item.setIcon(QIcon(pixmap))
                break

    def clear_files(self):
        self.clear()
        self.file_paths.clear()
        self._set_placeholder()
        self.files_changed.emit()

    def get_all_file_paths(self):
        paths = []
        for i in range(self.count()):
            item = self.item(i)
            if item != self._placeholder_item:
                paths.append(item.data(Qt.ItemDataRole.UserRole))
        return paths


# ──────────────────────────────────────────────────────────────────────────
#  Worker — semua logika konversi (pure PIL / pypdfium2 / pikepdf)
# ──────────────────────────────────────────────────────────────────────────
class PdfToolsWorker(QObject):
    progress_updated = Signal(int, str)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, mode, input_paths, output_path, **kwargs):
        super().__init__()
        self.mode = mode
        self.input_paths = input_paths
        self.output_path = output_path
        self.kwargs = kwargs
        self.is_running = True

    def stop(self):
        self.is_running = False

    @staticmethod
    def _parse_page_selection(sel_str, total_pages):
        if not sel_str.strip():
            return list(range(total_pages))
        pages = set()
        for part in sel_str.split(','):
            part = part.strip()
            if '-' in part:
                a, b = part.split('-', 1)
                try:
                    pages.update(range(int(a) - 1, int(b)))
                except ValueError:
                    pass
            else:
                try:
                    pages.add(int(part) - 1)
                except ValueError:
                    pass
        valid = sorted(p for p in pages if 0 <= p < total_pages)
        return valid if valid else list(range(total_pages))

    @staticmethod
    def _quality_to_jpeg(quality_idx):
        return [100, 95, 85, 75, 50][min(quality_idx, 4)]

    @staticmethod
    def _dpi_from_index(dpi_idx):
        return [72, 150, 200, 300, 600][min(dpi_idx, 4)]

    def run(self):
        try:
            if self.mode == 'img2pdf_single':
                self._images_to_single_pdf()
            elif self.mode == 'img2pdf_multi':
                self._images_to_multi_pdf()
            elif self.mode == 'pdf2img':
                self._pdf_to_images()
            elif self.mode == 'merge':
                self._merge_pdfs()
            elif self.mode == 'pdf2txt':
                self._pdf_to_txt()
            elif self.mode == 'pdf2docx':
                self._pdf_to_docx()
            elif self.mode == 'pdf2xlsx':
                self._pdf_to_xlsx()
        except Exception as e:
            if self.is_running:
                self.error.emit(f"Error: {e}")

    # ---- Image -> PDF ----
    def _images_to_single_pdf(self):
        quality = self._quality_to_jpeg(self.kwargs.get('quality_idx', 2))
        target_mb = self.kwargs.get('target_mb', 0.0)
        output_name = self.kwargs.get('output_name', 'output.pdf')
        output_file = os.path.join(self.output_path, output_name)

        total = len(self.input_paths)
        pil_images = []
        for i, path in enumerate(self.input_paths):
            if not self.is_running:
                return
            self.progress_updated.emit(int((i / total) * 85), f"Processing image {i+1}/{total}...")
            pil_images.append(Image.open(path).convert('RGB'))

        if not pil_images:
            self.error.emit("No images.")
            return

        self.progress_updated.emit(90, "Saving PDF...")
        first, rest = pil_images[0], pil_images[1:]
        save_kwargs = {'format': 'PDF', 'save_all': True, 'append_images': rest,
                        'quality': max(20, quality - 10) if target_mb > 0 else quality}
        first.save(output_file, **save_kwargs)

        if target_mb > 0:
            actual_mb = os.path.getsize(output_file) / (1024 * 1024)
            if actual_mb > target_mb:
                q2 = max(20, int(quality * (target_mb / actual_mb) * 0.9))
                compressed = []
                for img in pil_images:
                    buf = io.BytesIO()
                    img.save(buf, format='JPEG', quality=q2)
                    buf.seek(0)
                    compressed.append(Image.open(buf).convert('RGB'))
                compressed[0].save(output_file, format='PDF', save_all=True,
                                    append_images=compressed[1:], quality=q2)

        self.progress_updated.emit(100, "Done!")
        self.finished.emit(f"OK|PDF saved: {os.path.basename(output_file)}")

    def _images_to_multi_pdf(self):
        quality = self._quality_to_jpeg(self.kwargs.get('quality_idx', 2))
        target_mb = self.kwargs.get('target_mb', 0.0)
        total = len(self.input_paths)
        for i, path in enumerate(self.input_paths):
            if not self.is_running:
                return
            self.progress_updated.emit(int((i / total) * 100), f"Processing image {i+1}/{total}...")
            base = os.path.splitext(os.path.basename(path))[0]
            out_file = os.path.join(self.output_path, f"{base}.pdf")
            img = Image.open(path).convert('RGB')
            q = quality
            if target_mb > 0:
                buf = io.BytesIO()
                img.save(buf, format='PDF', quality=quality)
                actual_mb = buf.tell() / (1024 * 1024)
                if actual_mb > target_mb:
                    q = max(20, int(quality * (target_mb / actual_mb) * 0.85))
            img.save(out_file, format='PDF', quality=q)

        self.progress_updated.emit(100, "Done!")
        self.finished.emit(f"OK|{total} PDF files created.")

    # ---- PDF -> Image ----
    def _pdf_to_images(self):
        if not HAS_PDFIUM:
            self.error.emit("'pypdfium2' is not installed.")
            return
        fmt = self.kwargs.get('fmt', 'PNG').lower()
        dpi = self._dpi_from_index(self.kwargs.get('dpi_idx', 2))
        page_sel_str = self.kwargs.get('page_sel', '')
        scale = dpi / 72.0

        total_exported = 0
        for pdf_path in self.input_paths:
            if not self.is_running:
                return
            pdf = pdfium.PdfDocument(pdf_path)
            total_pages = len(pdf)
            pages = self._parse_page_selection(page_sel_str, total_pages)
            base = os.path.splitext(os.path.basename(pdf_path))[0]
            for j, page_idx in enumerate(pages):
                if not self.is_running:
                    return
                self.progress_updated.emit(
                    int((j / max(len(pages), 1)) * 100),
                    f"Converting page {j+1}/{len(pages)}...")
                pil_img = pdf[page_idx].render(scale=scale).to_pil()
                out_file = os.path.join(self.output_path, f"{base}_p{page_idx+1}.{fmt}")
                if fmt in ('jpg', 'jpeg', 'webp'):
                    pil_img.convert('RGB').save(out_file)
                else:
                    pil_img.save(out_file)
                total_exported += 1

        self.progress_updated.emit(100, "Done!")
        self.finished.emit(f"OK|{total_exported} images exported.")

    # ---- PDF Merger ----
    @staticmethod
    def _recompress_images_in_pdf(pdf_pikepdf, jpeg_quality):
        count = 0
        for page in pdf_pikepdf.pages:
            resources = page.get("/Resources")
            if resources is None:
                continue
            xobjects = resources.get("/XObject")
            if xobjects is None:
                continue
            for name in list(xobjects.keys()):
                xobj = xobjects[name]
                try:
                    if xobj.get("/Subtype") != "/Image":
                        continue
                    w = int(xobj["/Width"]); h = int(xobj["/Height"])
                    if w * h < 4096:
                        continue
                    cs = xobj.get("/ColorSpace")
                    if cs in ("/DeviceGray", "/DeviceCMYK"):
                        continue
                    raw = bytes(xobj.read_raw_bytes())
                    img = Image.open(io.BytesIO(raw))
                    if img.mode not in ("RGB", "L"):
                        img = img.convert("RGB")
                    buf = io.BytesIO()
                    img.save(buf, "JPEG", quality=jpeg_quality, optimize=True)
                    buf.seek(0)
                    xobj.write(buf.read(), filter=pikepdf.Name("/DCTDecode"), decode_parms=None)
                    xobj["/BitsPerComponent"] = 8
                    count += 1
                except Exception:
                    pass
        return count

    def _merge_pdfs(self):
        quality_idx = self.kwargs.get('quality_idx', 0)
        target_mb = self.kwargs.get('target_mb', 0.0)
        output_name = self.kwargs.get('output_name', 'merged.pdf')
        output_file = os.path.join(self.output_path, output_name)

        if not HAS_PIKEPDF:
            self._merge_pdfs_fallback(output_file)
            return

        jpeg_quality_map = {0: None, 1: 72, 2: 45}
        jpeg_quality = jpeg_quality_map.get(quality_idx, None)

        total_files = len(self.input_paths)
        self.progress_updated.emit(2, "Opening files...")
        try:
            merged = pikepdf.Pdf.new()
            for i, path in enumerate(self.input_paths):
                if not self.is_running:
                    return
                src = pikepdf.Pdf.open(path)
                merged.pages.extend(src.pages)
                self.progress_updated.emit(int(((i + 1) / total_files) * 50), f"Merging {i+1}/{total_files}...")
        except Exception as e:
            self.error.emit(f"Merge error: {e}")
            return

        if not self.is_running:
            return

        self.progress_updated.emit(55, "Optimizing...")
        try:
            merged.remove_unreferenced_resources()
        except Exception:
            pass

        if jpeg_quality is not None:
            self.progress_updated.emit(60, "Recompressing images...")
            try:
                self._recompress_images_in_pdf(merged, jpeg_quality)
            except Exception:
                pass

        self.progress_updated.emit(85, "Saving...")
        try:
            save_opts = dict(compress_streams=True,
                              object_stream_mode=pikepdf.ObjectStreamMode.generate,
                              recompress_flate=True)
            merged.save(output_file, **save_opts)

            if target_mb > 0:
                actual_mb = os.path.getsize(output_file) / (1024 * 1024)
                iteration = 0
                q = jpeg_quality if jpeg_quality else 65
                while actual_mb > target_mb and iteration < 3 and self.is_running:
                    iteration += 1
                    q = max(10, int(q * (target_mb / actual_mb) * 0.88))
                    self.progress_updated.emit(85 + iteration * 4, f"Shrinking (q={q})...")
                    try:
                        self._recompress_images_in_pdf(merged, q)
                        merged.save(output_file, **save_opts)
                        actual_mb = os.path.getsize(output_file) / (1024 * 1024)
                    except Exception:
                        break
        except Exception as e:
            self.error.emit(f"Save error: {e}")
            return

        self.progress_updated.emit(100, "Done!")
        self.finished.emit(f"OK|Merged PDF saved: {os.path.basename(output_file)}")

    def _merge_pdfs_fallback(self, output_file):
        if not HAS_PDFIUM:
            self.error.emit("'pypdfium2' or 'pikepdf' is not installed.")
            return
        merged = pdfium.PdfDocument.new()
        docs, total_pages = [], 0
        for path in self.input_paths:
            doc = pdfium.PdfDocument(path)
            docs.append(doc)
            total_pages += len(doc)
        done = 0
        for doc in docs:
            if not self.is_running:
                return
            merged.import_pages(doc, list(range(len(doc))))
            done += len(doc)
            self.progress_updated.emit(int((done / max(total_pages, 1)) * 95), "Merging...")
        merged.save(output_file)
        self.progress_updated.emit(100, "Done!")
        self.finished.emit(f"OK|Merged PDF saved: {os.path.basename(output_file)}")

    # ---- PDF Document Conversion ----
    def _pdf_to_txt(self):
        if not HAS_PDFIUM:
            self.error.emit("'pypdfium2' is not installed.")
            return
        total = len(self.input_paths)
        done_files = []
        for i, pdf_path in enumerate(self.input_paths):
            if not self.is_running:
                return
            self.progress_updated.emit(int((i / total) * 95), f"Extracting text {i+1}/{total}...")
            pdf = pdfium.PdfDocument(pdf_path)
            text_parts = []
            for page in pdf:
                textpage = page.get_textpage()
                text_parts.append(textpage.get_text_range())
            base = os.path.splitext(os.path.basename(pdf_path))[0]
            out_file = os.path.join(self.output_path, f"{base}.txt")
            with open(out_file, 'w', encoding='utf-8') as f:
                f.write("\n\n".join(text_parts))
            done_files.append(out_file)
        self.progress_updated.emit(100, "Done!")
        self.finished.emit(f"OK|{len(done_files)} TXT files created.")

    def _pdf_to_docx(self):
        if not HAS_PDFIUM:
            self.error.emit("'pypdfium2' is not installed.")
            return
        if not HAS_DOCX:
            self.error.emit("'python-docx' is not installed. Install: pip install python-docx --break-system-packages")
            return
        import docx
        total = len(self.input_paths)
        done_files = []
        for i, pdf_path in enumerate(self.input_paths):
            if not self.is_running:
                return
            self.progress_updated.emit(int((i / total) * 95), f"Converting {i+1}/{total}...")
            pdf = pdfium.PdfDocument(pdf_path)
            document = docx.Document()
            for p_idx, page in enumerate(pdf):
                textpage = page.get_textpage()
                text = textpage.get_text_range()
                for line in text.split('\n'):
                    if line.strip():
                        document.add_paragraph(line)
                if p_idx < len(pdf) - 1:
                    document.add_page_break()
            base = os.path.splitext(os.path.basename(pdf_path))[0]
            out_file = os.path.join(self.output_path, f"{base}.docx")
            document.save(out_file)
            done_files.append(out_file)
        self.progress_updated.emit(100, "Done!")
        self.finished.emit(f"OK|{len(done_files)} DOCX files created.")

    def _pdf_to_xlsx(self):
        if not HAS_PDFIUM:
            self.error.emit("'pypdfium2' is not installed.")
            return
        if not HAS_OPENPYXL:
            self.error.emit("'openpyxl' is not installed. Install: pip install openpyxl --break-system-packages")
            return
        import openpyxl
        total = len(self.input_paths)
        done_files = []
        for i, pdf_path in enumerate(self.input_paths):
            if not self.is_running:
                return
            self.progress_updated.emit(int((i / total) * 95), f"Converting {i+1}/{total}...")
            pdf = pdfium.PdfDocument(pdf_path)
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "PDF Text"
            row = 1
            for page in pdf:
                textpage = page.get_textpage()
                text = textpage.get_text_range()
                for line in text.split('\n'):
                    if line.strip():
                        ws.cell(row=row, column=1, value=line)
                        row += 1
            base = os.path.splitext(os.path.basename(pdf_path))[0]
            out_file = os.path.join(self.output_path, f"{base}.xlsx")
            wb.save(out_file)
            done_files.append(out_file)
        self.progress_updated.emit(100, "Done!")
        self.finished.emit(f"OK|{len(done_files)} XLSX files created.")


# ──────────────────────────────────────────────────────────────────────────
#  Output folder picker (reusable row widget)
# ──────────────────────────────────────────────────────────────────────────
class OutputFolderRow(QWidget):
    def __init__(self, lang, parent=None):
        super().__init__(parent)
        layout = QHBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        self.output_label = QLabel(lang["output_label"])
        self.path_edit = QLineEdit()
        self.path_edit.setReadOnly(True)
        self.path_edit.setPlaceholderText(lang["output_placeholder"])
        self.browse_btn = QPushButton(lang["choose_folder_btn"])
        self.browse_btn.clicked.connect(self._browse)
        layout.addWidget(self.output_label)
        layout.addWidget(self.path_edit, 1)
        layout.addWidget(self.browse_btn)
        self.lang = lang

    def retranslate(self, lang):
        self.lang = lang
        self.output_label.setText(lang["output_label"])
        self.path_edit.setPlaceholderText(lang["output_placeholder"])
        self.browse_btn.setText(lang["choose_folder_btn"])

    def _browse(self):
        folder = QFileDialog.getExistingDirectory(self, self.lang["choose_folder_title"])
        if folder:
            self.path_edit.setText(folder)

    def get_path(self):
        return self.path_edit.text()


# ──────────────────────────────────────────────────────────────────────────
#  Base page: drop area + output + progress + start/stop (shared scaffolding)
# ──────────────────────────────────────────────────────────────────────────
class BaseToolPage(QWidget):
    def __init__(self, title_key, accept_types, lang, parent=None):
        super().__init__(parent)
        self.thread = None
        self.worker = None
        self.lang = lang
        self.title_key = title_key

        root = QHBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setChildrenCollapsible(False)
        root.addWidget(splitter)

        # ── main area ──
        main_area = QWidget()
        main_layout = QVBoxLayout(main_area)
        main_layout.setContentsMargins(15, 15, 15, 15)
        main_layout.setSpacing(8)

        self.title_label = QLabel(f"<b>{lang[title_key]}</b>")
        self.title_label.setStyleSheet("font-size: 13pt; background: transparent;")
        main_layout.addWidget(self.title_label)

        btn_row = QHBoxLayout()
        self.add_btn = QPushButton(lang["add_files_btn"])
        self.add_btn.clicked.connect(self._browse_files)
        self.clear_btn = QPushButton(lang["clear_btn"])
        self.clear_btn.clicked.connect(self._clear_files)
        btn_row.addWidget(self.add_btn)
        btn_row.addWidget(self.clear_btn)
        btn_row.addStretch()
        main_layout.addLayout(btn_row)

        self.accept_types = accept_types
        self.drop_area = FileDropArea(accept_types=accept_types, lang=lang)
        main_layout.addWidget(self.drop_area, 1)

        self.output_row = OutputFolderRow(lang)
        main_layout.addWidget(self.output_row)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        main_layout.addWidget(self.progress_bar)

        self.status_label = QLabel(lang["status_ready"])
        self.status_label.setStyleSheet("background: transparent;")
        main_layout.addWidget(self.status_label)

        action_row = QHBoxLayout()
        self.start_btn = QPushButton(lang["start_btn"])
        self.start_btn.setObjectName("startButton")
        self.start_btn.clicked.connect(self._on_start_clicked)
        self.stop_btn = QPushButton(lang["stop_btn"])
        self.stop_btn.setEnabled(False)
        self.stop_btn.clicked.connect(self._on_stop_clicked)
        action_row.addWidget(self.start_btn)
        action_row.addWidget(self.stop_btn)
        main_layout.addLayout(action_row)

        splitter.addWidget(main_area)

        # ── options sidebar (filled by subclass) ──
        self.sidebar = QFrame()
        self.sidebar.setObjectName("optionsSidebar")
        self.sidebar.setMinimumWidth(260)
        self.sidebar_layout = QVBoxLayout(self.sidebar)
        self.sidebar_layout.setContentsMargins(12, 12, 12, 12)
        self.sidebar_layout.setSpacing(8)
        splitter.addWidget(self.sidebar)
        splitter.setSizes([600, 280])

        self.options_label = QLabel(f"<b>{lang['options_label']}</b>")
        self.options_label.setStyleSheet("background: transparent;")
        self.sidebar_layout.addWidget(self.options_label)

        self.build_options(self.sidebar_layout, lang)
        self.sidebar_layout.addStretch()

    # subclasses override this to add their option widgets
    def build_options(self, layout, lang):
        pass

    def retranslate_options(self, lang):
        pass

    def retranslate(self, lang):
        self.lang = lang
        self.title_label.setText(f"<b>{lang[self.title_key]}</b>")
        self.add_btn.setText(lang["add_files_btn"])
        self.clear_btn.setText(lang["clear_btn"])
        self.drop_area.retranslate(lang)
        self.output_row.retranslate(lang)
        self.status_label.setText(lang["status_ready"])
        self.stop_btn.setText(lang["stop_btn"])
        self.options_label.setText(f"<b>{lang['options_label']}</b>")
        self.retranslate_options(lang)

    def _browse_files(self):
        filters = []
        if 'image' in self.accept_types:
            filters.append("*.png *.jpg *.jpeg *.bmp *.webp *.gif")
        if 'pdf' in self.accept_types:
            filters.append("*.pdf")
        filter_str = "Files (" + " ".join(filters) + ")"
        files, _ = QFileDialog.getOpenFileNames(self, self.lang["add_files_btn"], "", filter_str)
        if files:
            self.drop_area.add_files(files)

    def _clear_files(self):
        self.drop_area.clear_files()

    def _filtered_paths(self, exts):
        return [p for p in self.drop_area.get_all_file_paths()
                if os.path.splitext(p)[1].lower() in exts]

    def _validate_output(self):
        out = self.output_row.get_path()
        if not out or not os.path.isdir(out):
            QMessageBox.warning(self, self.lang["error_title"], self.lang["error_no_output"])
            return None
        return out

    def _on_start_clicked(self):
        raise NotImplementedError

    def _run_worker(self, worker):
        self.thread = QThread()
        self.worker = worker
        self.worker.moveToThread(self.thread)
        self.thread.started.connect(self.worker.run)
        self.worker.progress_updated.connect(self._on_progress)
        self.worker.finished.connect(self._on_finished)
        self.worker.error.connect(self._on_error)
        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.thread.start()

    def _on_stop_clicked(self):
        if self.worker:
            self.worker.stop()
        self.status_label.setText(self.lang["status_stopped"])
        self._cleanup_thread()

    @Slot(int, str)
    def _on_progress(self, value, message):
        self.progress_bar.setValue(value)
        self.status_label.setText(message)

    @Slot(str)
    def _on_finished(self, message):
        # message format "OK|detail text"
        detail = message.split("|", 1)[1] if "|" in message else message
        self.progress_bar.setValue(100)
        self.status_label.setText(detail)
        QMessageBox.information(self, self.lang["done_title"], detail)
        self._open_output_folder()
        self._cleanup_thread()

    def _open_output_folder(self):
        out = self.output_row.get_path()
        try:
            open_in_file_manager(out)
        except Exception as e:
            QMessageBox.warning(self, self.lang["error_title"],
                                 self.lang["open_folder_failed"].format(err=e))

    @Slot(str)
    def _on_error(self, message):
        self.status_label.setText(message)
        QMessageBox.critical(self, self.lang["error_title"], message)
        self._cleanup_thread()

    def _cleanup_thread(self):
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        if self.thread:
            self.thread.quit()
            self.thread.wait()
            self.thread = None
        self.worker = None


# ──────────────────────────────────────────────────────────────────────────
#  4 sub-tool pages
# ──────────────────────────────────────────────────────────────────────────
class ImageToPdfPage(BaseToolPage):
    def __init__(self, lang, parent=None):
        super().__init__("img2pdf_title", ['image'], lang, parent)
        self.start_btn.setText(lang["img2pdf_start_btn"])

    def build_options(self, layout, lang):
        grid = QGridLayout(); grid.setSpacing(6)
        grid.setColumnMinimumWidth(0, 110); grid.setColumnStretch(1, 1)

        self.quality_label = QLabel(lang["quality_label"])
        grid.addWidget(self.quality_label, 0, 0)
        self.quality_combo = QComboBox()
        self.quality_combo.addItems(lang["qualities"])
        self.quality_combo.setCurrentIndex(2)
        grid.addWidget(self.quality_combo, 0, 1)

        self.out_mode_label = QLabel(lang["out_mode_label"])
        grid.addWidget(self.out_mode_label, 1, 0)
        self.output_combo = QComboBox()
        self.output_combo.addItems(lang["out_modes"])
        grid.addWidget(self.output_combo, 1, 1)

        self.target_size_label = QLabel(lang["target_size_label"])
        grid.addWidget(self.target_size_label, 2, 0)
        self.target_spin = QSpinBox()
        self.target_spin.setRange(0, 9999)
        self.target_spin.setSuffix(" MB")
        self.target_spin.setToolTip(lang["target_size_hint"])
        grid.addWidget(self.target_spin, 2, 1)

        self.estimated_size_label_title = QLabel(lang["estimated_size_label"])
        grid.addWidget(self.estimated_size_label_title, 3, 0)
        self.estimated_size_value = QLabel(lang["estimated_size_empty"])
        self.estimated_size_value.setStyleSheet("color: #9FBF9F; font-weight: bold; background: transparent;")
        grid.addWidget(self.estimated_size_value, 3, 1)

        layout.addLayout(grid)

        # update estimasi setiap kali file / opsi berubah
        self.drop_area.files_changed.connect(self._update_estimated_size)
        self.quality_combo.currentIndexChanged.connect(self._update_estimated_size)
        self.output_combo.currentIndexChanged.connect(self._update_estimated_size)
        self.target_spin.valueChanged.connect(self._update_estimated_size)
        self._update_estimated_size()

    # Faktor kompresi kasar relatif terhadap ukuran file gambar asli,
    # dipakai hanya sebagai estimasi (bukan hitungan presisi).
    _QUALITY_SIZE_FACTOR = {0: 1.0, 1: 0.85, 2: 0.6, 3: 0.45, 4: 0.25}

    def _update_estimated_size(self):
        img_paths = self._filtered_paths(IMAGE_EXT)
        if not img_paths:
            self.estimated_size_value.setText(self.lang["estimated_size_empty"])
            return
        try:
            total_bytes = sum(os.path.getsize(p) for p in img_paths if os.path.isfile(p))
        except OSError:
            total_bytes = 0
        factor = self._QUALITY_SIZE_FACTOR.get(self.quality_combo.currentIndex(), 0.6)
        estimated_bytes = total_bytes * factor

        target_mb = self.target_spin.value()
        if target_mb > 0:
            estimated_bytes = min(estimated_bytes, target_mb * 1024 * 1024)
            self.estimated_size_value.setText(
                self.lang["estimated_size_target"].format(target=human_size(target_mb * 1024 * 1024)))
        else:
            self.estimated_size_value.setText(
                self.lang["estimated_size_calc"].format(size=human_size(estimated_bytes)))

    def retranslate_options(self, lang):
        self.start_btn.setText(lang["img2pdf_start_btn"])
        self.quality_label.setText(lang["quality_label"])
        self.out_mode_label.setText(lang["out_mode_label"])
        self.target_size_label.setText(lang["target_size_label"])
        self.target_spin.setToolTip(lang["target_size_hint"])
        self.estimated_size_label_title.setText(lang["estimated_size_label"])
        idx_q, idx_o = self.quality_combo.currentIndex(), self.output_combo.currentIndex()
        self.quality_combo.clear(); self.quality_combo.addItems(lang["qualities"]); self.quality_combo.setCurrentIndex(idx_q)
        self.output_combo.clear(); self.output_combo.addItems(lang["out_modes"]); self.output_combo.setCurrentIndex(idx_o)
        self._update_estimated_size()

    def _on_start_clicked(self):
        img_paths = self._filtered_paths(IMAGE_EXT)
        if not img_paths:
            QMessageBox.warning(self, self.lang["error_title"], self.lang["img2pdf_no_files"])
            return
        out = self._validate_output()
        if not out:
            return
        single = self.output_combo.currentIndex() == 0
        mode = 'img2pdf_single' if single else 'img2pdf_multi'
        worker = PdfToolsWorker(
            mode=mode, input_paths=img_paths, output_path=out,
            quality_idx=self.quality_combo.currentIndex(),
            target_mb=float(self.target_spin.value()),
            output_name="images_combined.pdf" if single else "")
        self._run_worker(worker)


class PdfToImagePage(BaseToolPage):
    def __init__(self, lang, parent=None):
        super().__init__("pdf2img_title", ['pdf'], lang, parent)
        self.start_btn.setText(lang["pdf2img_start_btn"])

    def build_options(self, layout, lang):
        grid = QGridLayout(); grid.setSpacing(6)
        grid.setColumnMinimumWidth(0, 110); grid.setColumnStretch(1, 1)

        self.format_label = QLabel(lang["format_label"])
        grid.addWidget(self.format_label, 0, 0)
        self.format_combo = QComboBox()
        self.format_combo.addItems(["PNG", "JPG", "WEBP"])
        grid.addWidget(self.format_combo, 0, 1)

        self.dpi_label = QLabel(lang["dpi_label"])
        grid.addWidget(self.dpi_label, 1, 0)
        self.dpi_combo = QComboBox()
        self.dpi_combo.addItems(lang["dpis"])
        self.dpi_combo.setCurrentIndex(2)
        grid.addWidget(self.dpi_combo, 1, 1)

        self.pages_label = QLabel(lang["pages_label"])
        grid.addWidget(self.pages_label, 2, 0)
        self.pages_edit = QLineEdit()
        self.pages_edit.setPlaceholderText(lang["pages_hint"])
        grid.addWidget(self.pages_edit, 2, 1)

        layout.addLayout(grid)

    def retranslate_options(self, lang):
        self.start_btn.setText(lang["pdf2img_start_btn"])
        self.format_label.setText(lang["format_label"])
        self.dpi_label.setText(lang["dpi_label"])
        self.pages_label.setText(lang["pages_label"])
        self.pages_edit.setPlaceholderText(lang["pages_hint"])
        idx_d = self.dpi_combo.currentIndex()
        self.dpi_combo.clear(); self.dpi_combo.addItems(lang["dpis"]); self.dpi_combo.setCurrentIndex(idx_d)

    def _on_start_clicked(self):
        pdf_paths = self._filtered_paths(PDF_EXT)
        if not pdf_paths:
            QMessageBox.warning(self, self.lang["error_title"], self.lang["pdf2img_no_files"])
            return
        out = self._validate_output()
        if not out:
            return
        worker = PdfToolsWorker(
            mode='pdf2img', input_paths=pdf_paths, output_path=out,
            fmt=self.format_combo.currentText(),
            dpi_idx=self.dpi_combo.currentIndex(),
            page_sel=self.pages_edit.text())
        self._run_worker(worker)


class PdfMergerPage(BaseToolPage):
    def __init__(self, lang, parent=None):
        super().__init__("merger_title", ['pdf'], lang, parent)
        self.start_btn.setText(lang["merger_start_btn"])

    def build_options(self, layout, lang):
        grid = QGridLayout(); grid.setSpacing(6)
        grid.setColumnMinimumWidth(0, 110); grid.setColumnStretch(1, 1)

        self.quality_label = QLabel(lang["quality_label"])
        grid.addWidget(self.quality_label, 0, 0)
        self.quality_combo = QComboBox()
        self.quality_combo.addItems(lang["merger_qualities"])
        grid.addWidget(self.quality_combo, 0, 1)

        self.target_size_label = QLabel(lang["target_size_label"])
        grid.addWidget(self.target_size_label, 1, 0)
        self.target_spin = QSpinBox()
        self.target_spin.setRange(0, 9999)
        self.target_spin.setSuffix(" MB")
        self.target_spin.setToolTip(lang["target_size_hint"])
        grid.addWidget(self.target_spin, 1, 1)

        self.estimated_size_label_title = QLabel(lang["estimated_size_label"])
        grid.addWidget(self.estimated_size_label_title, 2, 0)
        self.estimated_size_value = QLabel(lang["estimated_size_empty"])
        self.estimated_size_value.setStyleSheet("color: #9FBF9F; font-weight: bold; background: transparent;")
        grid.addWidget(self.estimated_size_value, 2, 1)

        layout.addLayout(grid)

        self.pikepdf_note = None
        if not HAS_PIKEPDF:
            self.pikepdf_note = QLabel(lang["merger_pikepdf_note"])
            self.pikepdf_note.setStyleSheet("color: #d08770; font-size: 8pt; background: transparent;")
            layout.addWidget(self.pikepdf_note)

        self.drop_area.files_changed.connect(self._update_estimated_size)
        self.quality_combo.currentIndexChanged.connect(self._update_estimated_size)
        self.target_spin.valueChanged.connect(self._update_estimated_size)
        self._update_estimated_size()

    # Faktor kompresi kasar relatif terhadap total ukuran PDF asli,
    # dipakai hanya sebagai estimasi (bukan hitungan presisi).
    _MERGE_SIZE_FACTOR = {0: 1.0, 1: 0.55, 2: 0.35}

    def _update_estimated_size(self):
        pdf_paths = self._filtered_paths(PDF_EXT)
        if not pdf_paths:
            self.estimated_size_value.setText(self.lang["estimated_size_empty"])
            return
        try:
            total_bytes = sum(os.path.getsize(p) for p in pdf_paths if os.path.isfile(p))
        except OSError:
            total_bytes = 0
        factor = self._MERGE_SIZE_FACTOR.get(self.quality_combo.currentIndex(), 1.0)
        if not HAS_PIKEPDF:
            factor = 1.0  # tanpa pikepdf hanya merge struktural, tidak ada kompresi
        estimated_bytes = total_bytes * factor

        target_mb = self.target_spin.value()
        if target_mb > 0 and HAS_PIKEPDF:
            estimated_bytes = min(estimated_bytes, target_mb * 1024 * 1024)
            self.estimated_size_value.setText(
                self.lang["estimated_size_target"].format(target=human_size(target_mb * 1024 * 1024)))
        else:
            self.estimated_size_value.setText(
                self.lang["estimated_size_calc"].format(size=human_size(estimated_bytes)))

    def retranslate_options(self, lang):
        self.start_btn.setText(lang["merger_start_btn"])
        self.quality_label.setText(lang["quality_label"])
        self.target_size_label.setText(lang["target_size_label"])
        self.target_spin.setToolTip(lang["target_size_hint"])
        self.estimated_size_label_title.setText(lang["estimated_size_label"])
        idx_q = self.quality_combo.currentIndex()
        self.quality_combo.clear(); self.quality_combo.addItems(lang["merger_qualities"]); self.quality_combo.setCurrentIndex(idx_q)
        if self.pikepdf_note:
            self.pikepdf_note.setText(lang["merger_pikepdf_note"])
        self._update_estimated_size()

    def _on_start_clicked(self):
        pdf_paths = self._filtered_paths(PDF_EXT)
        if len(pdf_paths) < 2:
            QMessageBox.warning(self, self.lang["error_title"], self.lang["merger_no_files"])
            return
        out = self._validate_output()
        if not out:
            return
        worker = PdfToolsWorker(
            mode='merge', input_paths=pdf_paths, output_path=out,
            quality_idx=self.quality_combo.currentIndex(),
            target_mb=float(self.target_spin.value()),
            output_name="merged.pdf")
        self._run_worker(worker)


class PdfDocConversionPage(BaseToolPage):
    def __init__(self, lang, parent=None):
        super().__init__("nav_docconv", ['pdf'], lang, parent)
        self.start_btn.setText(lang["docconv_start_btn"])

    def build_options(self, layout, lang):
        grid = QGridLayout(); grid.setSpacing(6)
        grid.setColumnMinimumWidth(0, 110); grid.setColumnStretch(1, 1)

        self.convert_to_label = QLabel(lang["convert_to_label"])
        grid.addWidget(self.convert_to_label, 0, 0)
        self.target_combo = QComboBox()
        items = ["TXT"]
        if HAS_DOCX:
            items.append("DOCX")
        if HAS_OPENPYXL:
            items.append("XLSX")
        self.target_combo.addItems(items)
        grid.addWidget(self.target_combo, 0, 1)

        layout.addLayout(grid)

        self.missing_note = None
        missing = []
        if not HAS_DOCX:
            missing.append("python-docx (DOCX)")
        if not HAS_OPENPYXL:
            missing.append("openpyxl (XLSX)")
        if missing:
            self.missing_note = QLabel(lang["missing_libs_label"].format(items=", ".join(missing)))
            self.missing_note.setStyleSheet("color: #d08770; font-size: 8pt; background: transparent;")
            self.missing_note.setWordWrap(True)
            layout.addWidget(self.missing_note)
        self._missing_items = missing

    def retranslate_options(self, lang):
        self.start_btn.setText(lang["docconv_start_btn"])
        self.convert_to_label.setText(lang["convert_to_label"])
        if self.missing_note:
            self.missing_note.setText(lang["missing_libs_label"].format(items=", ".join(self._missing_items)))

    def _on_start_clicked(self):
        pdf_paths = self._filtered_paths(PDF_EXT)
        if not pdf_paths:
            QMessageBox.warning(self, self.lang["error_title"], self.lang["docconv_no_files"])
            return
        out = self._validate_output()
        if not out:
            return
        target = self.target_combo.currentText()
        mode_map = {"TXT": "pdf2txt", "DOCX": "pdf2docx", "XLSX": "pdf2xlsx"}
        worker = PdfToolsWorker(mode=mode_map[target], input_paths=pdf_paths, output_path=out)
        self._run_worker(worker)


# ──────────────────────────────────────────────────────────────────────────
#  Update checker — fetch version.json di background (tidak memblokir UI)
# ──────────────────────────────────────────────────────────────────────────
class _UpdateCheckSignals(QObject):
    finished = Signal(dict)
    error = Signal(str)


class UpdateCheckWorker(QRunnable):
    def __init__(self, url):
        super().__init__()
        self.url = url
        self.signals = _UpdateCheckSignals()

    def run(self):
        try:
            req = urllib.request.Request(self.url, headers={"User-Agent": "MacanPdfTools/1.0"})
            with urllib.request.urlopen(req, timeout=10) as resp:
                raw = resp.read().decode("utf-8")
            data = json.loads(raw)
            if not isinstance(data, dict):
                raise ValueError("Unexpected JSON format")
            self.signals.finished.emit(data)
        except urllib.error.URLError as e:
            self.signals.error.emit(str(getattr(e, "reason", e)))
        except Exception as e:
            self.signals.error.emit(str(e))


def _version_tuple(v):
    """Ubah string versi (mis. '1.3.0') jadi tuple int untuk dibandingkan."""
    parts = []
    for chunk in str(v).strip().split('.'):
        digits = ''.join(ch for ch in chunk if ch.isdigit())
        parts.append(int(digits) if digits else 0)
    return tuple(parts) if parts else (0,)


# ──────────────────────────────────────────────────────────────────────────
#  About page — info aplikasi, status dependency, dan cek pembaruan
# ──────────────────────────────────────────────────────────────────────────
class AboutPage(QWidget):
    def __init__(self, lang, parent=None):
        super().__init__(parent)
        self.lang = lang
        self.thread_pool = QThreadPool.globalInstance()

        outer = QVBoxLayout(self)
        outer.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        outer.addWidget(scroll)

        content = QWidget()
        scroll.setWidget(content)
        layout = QVBoxLayout(content)
        layout.setContentsMargins(30, 26, 30, 26)
        layout.setSpacing(18)

        # ── Header: logo + nama + versi ──
        header_row = QHBoxLayout()
        header_row.setSpacing(16)

        self.logo_label = QLabel()
        pix = QPixmap(LOGO_PATH)
        if not pix.isNull():
            pix = pix.scaledToHeight(64, Qt.TransformationMode.SmoothTransformation)
            self.logo_label.setPixmap(pix)
        self.logo_label.setFixedSize(64, 64)
        self.logo_label.setStyleSheet("background: transparent;")
        header_row.addWidget(self.logo_label)

        title_col = QVBoxLayout()
        title_col.setSpacing(2)
        self.app_title_label = QLabel(f"<span style='font-size:18pt; font-weight:700;'>{lang['about_app_title']}</span>")
        self.app_title_label.setStyleSheet("background: transparent;")
        self.edition_label = QLabel(lang["about_edition"])
        self.edition_label.setStyleSheet("color: #9AA0A6; background: transparent;")
        self.version_label = QLabel(f"{lang['about_version_label']} {APP_VERSION}")
        self.version_label.setStyleSheet("color: #C8C8C8; background: transparent;")
        title_col.addWidget(self.app_title_label)
        title_col.addWidget(self.edition_label)
        title_col.addWidget(self.version_label)
        header_row.addLayout(title_col)
        header_row.addStretch()
        layout.addLayout(header_row)

        self.tagline_label = QLabel(lang["about_tagline"])
        self.tagline_label.setWordWrap(True)
        self.tagline_label.setStyleSheet("color: #A3BE8C; font-style: italic; background: transparent;")
        layout.addWidget(self.tagline_label)

        layout.addWidget(self._hline())

        # ── Description ──
        self.description_label = QLabel(lang["about_description"])
        self.description_label.setWordWrap(True)
        self.description_label.setStyleSheet("background: transparent; line-height: 140%;")
        layout.addWidget(self.description_label)

        # ── Features ──
        self.features_title_label = QLabel(f"<b>{lang['about_features_title']}</b>")
        self.features_title_label.setStyleSheet("background: transparent; font-size: 11pt;")
        layout.addWidget(self.features_title_label)

        self.features_box = QVBoxLayout()
        self.features_box.setSpacing(4)
        layout.addLayout(self.features_box)
        self._populate_features(lang)

        layout.addWidget(self._hline())

        # ── Dependency status ──
        self.dep_title_label = QLabel(f"<b>{lang['about_dependencies_title']}</b>")
        self.dep_title_label.setStyleSheet("background: transparent; font-size: 11pt;")
        layout.addWidget(self.dep_title_label)

        self.dep_grid = QGridLayout()
        self.dep_grid.setSpacing(6)
        layout.addLayout(self.dep_grid)
        self._populate_dependencies(lang)

        layout.addWidget(self._hline())

        # ── Update check ──
        self.update_title_label = QLabel(f"<b>{lang['about_update_title']}</b>")
        self.update_title_label.setStyleSheet("background: transparent; font-size: 11pt;")
        layout.addWidget(self.update_title_label)

        update_row = QHBoxLayout()
        self.check_update_btn = QPushButton(lang["about_check_update_btn"])
        self.check_update_btn.setObjectName("startButton")
        self.check_update_btn.clicked.connect(self._check_for_update)
        update_row.addWidget(self.check_update_btn)

        self.open_download_btn = QPushButton(lang["about_open_download_btn"])
        self.open_download_btn.setVisible(False)
        self.open_download_btn.clicked.connect(self._open_download_url)
        update_row.addWidget(self.open_download_btn)
        update_row.addStretch()
        layout.addLayout(update_row)

        self.update_status_label = QLabel("")
        self.update_status_label.setWordWrap(True)
        self.update_status_label.setStyleSheet("background: transparent;")
        layout.addWidget(self.update_status_label)

        self.update_notes_title_label = QLabel(f"<b>{lang['about_update_notes_title']}</b>")
        self.update_notes_title_label.setStyleSheet("background: transparent;")
        self.update_notes_title_label.setVisible(False)
        layout.addWidget(self.update_notes_title_label)

        self.update_notes_label = QLabel("")
        self.update_notes_label.setWordWrap(True)
        self.update_notes_label.setStyleSheet("color: #C8C8C8; background: transparent;")
        layout.addWidget(self.update_notes_label)

        self._latest_download_url = None

        layout.addWidget(self._hline())

        # ── Links + credits ──
        link_row = QHBoxLayout()
        self.repo_btn = QPushButton(lang["about_open_repo_btn"])
        self.repo_btn.clicked.connect(lambda: QDesktopServices.openUrl(QUrl(REPO_URL)))
        link_row.addWidget(self.repo_btn)
        link_row.addStretch()
        layout.addLayout(link_row)

        self.credits_label = QLabel(lang["about_credits"])
        self.credits_label.setWordWrap(True)
        self.credits_label.setStyleSheet("color: #9AA0A6; font-size: 9pt; background: transparent;")
        layout.addWidget(self.credits_label)

        import datetime
        self.copyright_label = QLabel(lang["about_copyright"].format(year=datetime.date.today().year))
        self.copyright_label.setStyleSheet("color: #777777; font-size: 8pt; background: transparent;")
        layout.addWidget(self.copyright_label)

        layout.addStretch()

    def _hline(self):
        line = QFrame()
        line.setFrameShape(QFrame.Shape.HLine)
        line.setStyleSheet("color: #444444; background-color: #444444; max-height: 1px;")
        return line

    def _populate_features(self, lang):
        # bersihkan dulu kalau retranslate
        while self.features_box.count():
            item = self.features_box.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()
        for feat in lang["about_features"]:
            row = QLabel(f"•  {feat}")
            row.setWordWrap(True)
            row.setStyleSheet("background: transparent;")
            self.features_box.addWidget(row)

    def _populate_dependencies(self, lang):
        while self.dep_grid.count():
            item = self.dep_grid.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()
        deps = [
            ("pypdfium2", HAS_PDFIUM, True),
            ("pikepdf", HAS_PIKEPDF, False),
            ("python-docx", HAS_DOCX, False),
            ("openpyxl", HAS_OPENPYXL, False),
        ]
        for row, (name, installed, required) in enumerate(deps):
            tag = lang["about_dep_required"] if required else lang["about_dep_optional"]
            name_lbl = QLabel(f"{name}  <span style='color:#888888; font-size:8pt;'>{tag}</span>")
            name_lbl.setStyleSheet("background: transparent;")
            self.dep_grid.addWidget(name_lbl, row, 0)

            status_text = lang["about_dep_installed"] if installed else lang["about_dep_missing"]
            color = "#A3BE8C" if installed else ("#BF6160" if required else "#d08770")
            status_lbl = QLabel(f"●  {status_text}")
            status_lbl.setStyleSheet(f"color: {color}; background: transparent;")
            self.dep_grid.addWidget(status_lbl, row, 1)

    def retranslate(self, lang):
        self.lang = lang
        self.app_title_label.setText(f"<span style='font-size:18pt; font-weight:700;'>{lang['about_app_title']}</span>")
        self.edition_label.setText(lang["about_edition"])
        self.version_label.setText(f"{lang['about_version_label']} {APP_VERSION}")
        self.tagline_label.setText(lang["about_tagline"])
        self.description_label.setText(lang["about_description"])
        self.features_title_label.setText(f"<b>{lang['about_features_title']}</b>")
        self._populate_features(lang)
        self.dep_title_label.setText(f"<b>{lang['about_dependencies_title']}</b>")
        self._populate_dependencies(lang)
        self.update_title_label.setText(f"<b>{lang['about_update_title']}</b>")
        self.check_update_btn.setText(lang["about_check_update_btn"])
        self.open_download_btn.setText(lang["about_open_download_btn"])
        self.update_notes_title_label.setText(f"<b>{lang['about_update_notes_title']}</b>")
        self.repo_btn.setText(lang["about_open_repo_btn"])
        self.credits_label.setText(lang["about_credits"])
        import datetime
        self.copyright_label.setText(lang["about_copyright"].format(year=datetime.date.today().year))
        self.update_status_label.setText("")

    def _check_for_update(self):
        self.check_update_btn.setEnabled(False)
        self.open_download_btn.setVisible(False)
        self.update_notes_title_label.setVisible(False)
        self.update_notes_label.setText("")
        self.update_status_label.setStyleSheet("background: transparent; color: #C8C8C8;")
        self.update_status_label.setText(self.lang["about_checking_update"])

        worker = UpdateCheckWorker(UPDATE_JSON_URL)
        worker.signals.finished.connect(self._on_update_finished)
        worker.signals.error.connect(self._on_update_error)
        self.thread_pool.start(worker)

    @Slot(dict)
    def _on_update_finished(self, data):
        self.check_update_btn.setEnabled(True)
        latest_version = str(data.get("version", "")).strip()
        notes = data.get("notes") or data.get("changelog") or data.get("release_notes") or ""
        download_url = data.get("release_url") or data.get("download_url") or data.get("url") or REPO_URL
        self._latest_download_url = download_url

        if not latest_version:
            self._on_update_error("Invalid version.json (missing 'version' field).")
            return

        if _version_tuple(latest_version) > _version_tuple(APP_VERSION):
            self.update_status_label.setStyleSheet("background: transparent; color: #A3BE8C; font-weight: bold;")
            self.update_status_label.setText(
                self.lang["about_update_available"].format(version=latest_version, current=APP_VERSION))
            self.open_download_btn.setVisible(True)
            if notes:
                self.update_notes_title_label.setVisible(True)
                self.update_notes_label.setText(str(notes))
        else:
            self.update_status_label.setStyleSheet("background: transparent; color: #9AA0A6;")
            self.update_status_label.setText(self.lang["about_up_to_date"].format(version=APP_VERSION))

    @Slot(str)
    def _on_update_error(self, err):
        self.check_update_btn.setEnabled(True)
        self.update_status_label.setStyleSheet("background: transparent; color: #BF6160;")
        self.update_status_label.setText(self.lang["about_update_error"].format(err=err))

    def _open_download_url(self):
        if self._latest_download_url:
            QDesktopServices.openUrl(QUrl(self._latest_download_url))


# ──────────────────────────────────────────────────────────────────────────
#  Main window
# ──────────────────────────────────────────────────────────────────────────
class MacanPdfToolsApp(QMainWindow):
    def __init__(self):
        super().__init__()
        icon_path = "icon_pdf.ico"
        if hasattr(sys, "_MEIPASS"):
            icon_path = os.path.join(sys._MEIPASS, icon_path)
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        self.settings = QSettings(ORG_NAME, APP_NAME)

        # ── Muat pengaturan tersimpan (bahasa, geometri, posisi window) ──
        saved_lang = self.settings.value("language", "id")
        self.current_lang_code = saved_lang if saved_lang in LANGUAGES else "id"
        self.lang = LANGUAGES[self.current_lang_code]

        central = QWidget()
        self.setCentralWidget(central)
        root = QVBoxLayout(central)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # ── top bar: language switcher ──
        top_bar = QWidget()
        top_bar.setObjectName("topBar")
        top_layout = QHBoxLayout(top_bar)
        top_layout.setContentsMargins(12, 6, 12, 6)
        self.lang_label = QLabel(self.lang["lang_label"])
        self.lang_label.setStyleSheet("background: transparent;")
        top_layout.addWidget(self.lang_label)
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(["Bahasa Indonesia", "English"])
        self.lang_combo.setFixedWidth(160)
        self.lang_combo.setCurrentIndex(0 if self.current_lang_code == "id" else 1)
        self.lang_combo.currentIndexChanged.connect(self._on_lang_changed)
        top_layout.addWidget(self.lang_combo)
        top_layout.addStretch()
        root.addWidget(top_bar)

        body = QWidget()
        body_layout = QHBoxLayout(body)
        body_layout.setContentsMargins(0, 0, 0, 0)
        body_layout.setSpacing(0)
        root.addWidget(body, 1)

        # left nav
        self.nav_list = QListWidget()
        self.nav_list.setObjectName("navList")
        self.nav_list.setFixedWidth(210)
        self.nav_list.setIconSize(QSize(20, 20))
        self._populate_nav()
        self.nav_list.currentRowChanged.connect(self._on_nav_changed)
        body_layout.addWidget(self.nav_list)

        # stacked pages
        self.stack = QStackedWidget()
        self.pages = [
            ImageToPdfPage(self.lang),
            PdfToImagePage(self.lang),
            PdfMergerPage(self.lang),
            PdfDocConversionPage(self.lang),
            AboutPage(self.lang),
        ]
        for p in self.pages:
            self.stack.addWidget(p)
        body_layout.addWidget(self.stack, 1)

        self.nav_list.setCurrentRow(0)

        # ── footer panel dengan logo.png ──
        self.footer = self._build_footer()
        root.addWidget(self.footer)

        self._apply_stylesheet()
        self._set_window_title()
        self._check_dependencies()
        self._restore_window_geometry()

    def _build_footer(self):
        footer = QWidget()
        footer.setObjectName("footerPanel")
        footer.setFixedHeight(44)
        layout = QHBoxLayout(footer)
        layout.setContentsMargins(12, 4, 12, 4)
        layout.setSpacing(8)

        self.footer_logo_label = QLabel()
        self.footer_logo_label.setStyleSheet("background: transparent;")
        self._load_footer_logo()
        layout.addWidget(self.footer_logo_label)

        self.footer_text_label = QLabel(f"Macan Angkasa - All rights reserved")
        self.footer_text_label.setStyleSheet("background: transparent; color: #999999; font-size: 8pt;")
        layout.addWidget(self.footer_text_label)

        layout.addStretch()
        return footer

    def _load_footer_logo(self):
        pix = QPixmap(LOGO_PATH)
        if not pix.isNull():
            pix = pix.scaledToHeight(32, Qt.TransformationMode.SmoothTransformation)
            self.footer_logo_label.setPixmap(pix)
            self.footer_logo_label.setFixedSize(pix.size())
        else:
            # logo.png belum ada di folder app — sembunyikan area logo, jangan error
            self.footer_logo_label.setFixedSize(0, 0)

    def _restore_window_geometry(self):
        geometry = self.settings.value("window/geometry")
        if geometry is not None:
            self.restoreGeometry(geometry)
        else:
            self.resize(1100, 760)
            pos = self.settings.value("window/pos")
            if pos is not None:
                self.move(pos)

    def closeEvent(self, event):
        # Simpan bahasa, geometri (ukuran+state), dan posisi window
        self.settings.setValue("language", self.current_lang_code)
        self.settings.setValue("window/geometry", self.saveGeometry())
        self.settings.setValue("window/pos", self.pos())
        super().closeEvent(event)

    def _populate_nav(self):
        self.nav_list.clear()
        nav_items = [
            ("img2pdf", self.lang["nav_img2pdf"]),
            ("pdf2img", self.lang["nav_pdf2img"]),
            ("merger", self.lang["nav_merger"]),
            ("docconv", self.lang["nav_docconv"]),
            ("about", self.lang["nav_about"]),
        ]
        for icon_key, label in nav_items:
            item = QListWidgetItem(svg_to_icon(icon_key), label)
            self.nav_list.addItem(item)

    def _set_window_title(self):
        self.setWindowTitle(self.lang["window_title"].format(version=APP_VERSION))

    def _on_nav_changed(self, index):
        if index >= 0:
            self.stack.setCurrentIndex(index)

    def _on_lang_changed(self, index):
        self.current_lang_code = "id" if index == 0 else "en"
        self.lang = LANGUAGES[self.current_lang_code]
        self.settings.setValue("language", self.current_lang_code)

        current_row = self.nav_list.currentRow()
        self.lang_label.setText(self.lang["lang_label"])
        self._populate_nav()
        self.nav_list.setCurrentRow(current_row)
        self._set_window_title()

        for p in self.pages:
            p.retranslate(self.lang)

    def _check_dependencies(self):
        missing = []
        if not HAS_PDFIUM:
            missing.append("pypdfium2")
        if missing:
            QMessageBox.warning(
                self, self.lang["dep_missing_title"],
                self.lang["dep_missing_body"].format(items="\n".join(missing), pkgs=" ".join(missing)))

    def _apply_stylesheet(self):
        self.setStyleSheet("""
            QMainWindow, QWidget {
                background-color: #252525;
                color: #E0E0E0;
                font-family: Segoe UI, sans-serif;
            }
            QLabel {
                font-size: 10pt;
                background: transparent;
            }
            #topBar {
                background-color: #2c2c2c;
                border-bottom: 1px solid #444444;
            }
            #footerPanel {
                background-color: #2c2c2c;
                border-top: 1px solid #444444;
            }
            #navList {
                background-color: #333333;
                border: none;
                border-right: 1px solid #444444;
            }
            #navList::item {
                padding: 12px 10px;
                border-bottom: 1px solid #3a3a3a;
                background: transparent;
            }
            #navList::item:selected {
                background-color: #5A5A5A;
                color: #FFFFFF;
            }
            #navList::item:hover:!selected {
                background-color: #4A4A4A;
            }
            FileDropArea {
                background-color: #333333;
                border: 2px dashed #555555;
                border-radius: 5px;
                color: #AAAAAA;
            }
            FileDropArea::item {
                background-color: #404040;
                border: 1px solid #555555;
                border-radius: 4px;
                padding: 5px;
            }
            FileDropArea::item:selected {
                background-color: #5A5A5A;
                border-color: #A3BE8C;
            }
            #optionsSidebar {
                background-color: #2c2c2c;
                border-left: 1px solid #444444;
            }
            QLineEdit, QComboBox, QSpinBox {
                background-color: #3A3A3A;
                border: 1px solid #555555;
                border-radius: 4px;
                padding: 4px;
                color: #E0E0E0;
            }
            QPushButton {
                background-color: #4A4A4A;
                border: 1px solid #5A5A5A;
                border-radius: 4px;
                padding: 6px 12px;
            }
            QPushButton:hover { background-color: #5A5A5A; }
            QPushButton:disabled { color: #777777; background-color: #3a3a3a; }
            #startButton {
                background-color: #4C7A4C;
                font-weight: bold;
            }
            #startButton:hover { background-color: #5C8A5C; }
            QProgressBar {
                border: 1px solid #555555;
                border-radius: 4px;
                text-align: center;
                background-color: #333333;
            }
            QProgressBar::chunk { background-color: #5A8A5A; }
        """)


def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    window = MacanPdfToolsApp()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
